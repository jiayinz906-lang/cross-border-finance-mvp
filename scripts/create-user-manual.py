from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "XJD_Finance_系统使用操作说明书.docx"

NAVY = "17264A"
BLUE = "2F6FED"
TEAL = "20A67A"
ORANGE = "F08A24"
INK = "13213C"
MUTED = "67758D"
LIGHT_BLUE = "E8EEF5"
LIGHTER_BLUE = "F4F7FB"
LIGHT_GREEN = "EAF7F2"
LIGHT_ORANGE = "FFF4E6"
WHITE = "FFFFFF"
GRID = "C8D2E1"
RED = "B42318"

LATIN_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=LATIN_FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.25, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next


def add_text(doc, text: str, *, bold=False, color=INK, size=11, after=6, align=None):
    p = doc.add_paragraph()
    style_paragraph(p, after=after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_list(doc, items: list[str], numbered=False):
    style_name = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style_name)
        style_paragraph(p, after=4, line=1.25)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        run = p.add_run(item)
        set_run_font(run, size=11, color=INK)


def add_callout(doc, title: str, body: str, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "20")
    start.set(qn("w:color"), accent)
    borders.append(start)
    p = cell.paragraphs[0]
    style_paragraph(p, after=3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    style_paragraph(p2, after=0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_text(doc, "", size=1, after=2)


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int], header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, after=0, line=1.0)
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(values):
            if row_index % 2 == 1:
                set_cell_shading(cells[col_index], "F9FBFD")
            p = cells[col_index].paragraphs[0]
            if headers[col_index] == "序":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.1)
            r = p.add_run(str(value))
            set_run_font(r, size=9.2, color=INK)
    set_table_geometry(table, widths_dxa)
    add_text(doc, "", size=1, after=4)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: NAVY}[level], bold=True)
    p.paragraph_format.keep_with_next = True
    return p


def add_step(doc, number: int, title: str, body: str):
    p = doc.add_paragraph()
    style_paragraph(p, before=2, after=3, keep_with_next=True)
    r1 = p.add_run(f"步骤 {number}  {title}")
    set_run_font(r1, size=11.5, color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    style_paragraph(p2, after=7)
    p2.paragraph_format.left_indent = Inches(0.18)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, NAVY, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, after=0)
    r = p.add_run("XJD Finance  |  跨境物流财务管理使用操作说明书")
    set_run_font(r, size=9, color=MUTED, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7200, 2160], indent_dxa=0)
    table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = table.rows[0].cells[0].paragraphs[0].add_run("内部操作参考 | 数据以已导入原始 Excel 和数据库记录为准")
    set_run_font(r1, size=8.5, color=MUTED)
    p2 = table.rows[0].cells[1].paragraphs[0]
    r2 = p2.add_run("第 ")
    set_run_font(r2, size=9, color=MUTED)
    add_page_number(p2)
    r3 = p2.add_run(" 页")
    set_run_font(r3, size=9, color=MUTED)


def build_manual():
    doc = Document()
    configure_document(doc)

    # Cover: editorial_cover pattern with the compact_reference_guide token system.
    add_text(doc, "XJD FINANCE", bold=True, color=BLUE, size=12, after=54, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "系统使用操作说明书", bold=True, color=NAVY, size=30, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "跨境物流财务管理 · 月度数据 · 提成签名 · 应收应付", color=MUTED, size=14, after=38, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "适用于财务、主管、管理层、销售代表与操作员", color=INK, size=11, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"版本 1.0  |  {date.today().isoformat()}", color=MUTED, size=10, after=42, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(
        doc,
        "使用原则",
        "系统汇总以当前月份最新有效导入批次为准。应收、应付、毛利、赔付、汇率和人员归属均追溯到原始 Excel 明细行；发现差异时先核对原始数据，不直接修改汇总金额。",
        fill=LIGHTER_BLUE,
        accent=BLUE,
    )
    add_text(doc, "线上系统", bold=True, color=NAVY, size=10.5, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "https://jiayinz906-lang.github.io/cross-border-finance-mvp/#/dashboard", color=BLUE, size=9.5, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "本地开发：http://localhost:5173/  |  API：http://localhost:4000/api", color=MUTED, size=9, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_heading(doc, "手册概览", 1)
    add_text(doc, "本手册按真实月度业务处理顺序编写。普通用户从“快速开始”进入；主管和财务人员应重点阅读 Excel 导入、确认签名、应收应付和月结锁账章节。")
    add_table(
        doc,
        ["章节", "适用人员", "目标"],
        [
            ["快速开始与登录", "全部用户", "进入系统、选择月份并确认页面状态"],
            ["Excel 导入", "财务、主管、管理员", "预检原始台账并写入对应月份"],
            ["提成、绩效与签名", "主管、销售、操作员", "完成金额确认和证据链留存"],
            ["应收、应付与风险", "财务、主管", "登记收付款、复核异常并追溯原始行"],
            ["参数、月结与运维", "管理员、主管", "管理规则、锁账、账号、日志与服务状态"],
        ],
        [2100, 2400, 4860],
    )
    add_heading(doc, "1. 访问、账号与登录", 1)
    add_table(
        doc,
        ["环境", "地址", "说明"],
        [
            ["线上前端", "GitHub Pages / #/dashboard", "电脑关闭后仍可访问；用于正式操作"],
            ["线上后端", "Render /api", "由前端调用，不作为日常操作页面"],
            ["本地前端", "http://localhost:5173/", "仅本机开发测试，依赖本机服务"],
            ["本地后端", "http://localhost:4000/api", "健康检查路径为 /health"],
        ],
        [1800, 3900, 3660],
    )
    add_step(doc, 1, "打开系统", "正式使用请打开线上前端地址；本机地址只用于开发和排查。未登录访问业务页面会自动跳转到登录页。")
    add_step(doc, 2, "输入账号与密码", "使用管理员分配的独立账号。首次登录或管理员重置密码后，按页面提示立即修改密码。")
    add_step(doc, 3, "确认身份与月份", "左侧栏底部显示当前用户和角色，顶部或侧栏显示当前月份。操作前先确认月份正确。")
    add_callout(doc, "安全要求", "不要在聊天、文档或群消息中传播生产密码。默认初始化口令仅用于首次建库，正式使用前必须修改；员工离职后应立即停用账号。", fill=LIGHT_ORANGE, accent=ORANGE)

    add_heading(doc, "2. 角色与权限", 1)
    add_table(
        doc,
        ["角色", "主要可执行操作", "限制"],
        [
            ["系统管理员", "账号、导入、回滚、规则、风险、确认单、锁账、报表", "高权限操作均写入审计日志"],
            ["财务", "Excel 导入、风险复核、应收应付登记、报表导出", "不能修改参数规则或主管确认"],
            ["主管", "导入、风险复核、提成确认、签名主管确认、锁账", "不能管理系统账号或重置全部数据"],
            ["老板/管理层", "查看经营与财务数据、导出报表", "只读，不进行业务写入"],
            ["销售/客服", "查看业务、确认单和个人相关数据", "不能导入、锁账或修改规则"],
        ],
        [1800, 4680, 2880],
    )
    doc.add_page_break()

    add_heading(doc, "3. 月度工作快速流程", 1)
    add_step(doc, 1, "切换或确认月份", "在经营总览右上方点击“月份”，从数据库已有月份列表中选择。页面切换后，各业务页统一读取该月份最新有效批次。")
    add_step(doc, 2, "导入 Excel", "上传符合固定表头规范的月度运单明细，先完成预检，再确认写入数据库。")
    add_step(doc, 3, "核对汇总与风险", "检查总应收、总应付、毛利、票数、物流/服务拆分；进入风险复查核对低毛利、异常高利润或缺失信息。")
    add_step(doc, 4, "确认提成与绩效", "主管确认注册/服务提成，核对物流提成比例，确认操作员绩效规则与金额。")
    add_step(doc, 5, "生成并发送确认单", "在电子签名确认页生成销售薪资确认单和操作员薪资确认单，通过钉钉或复制外部链接发送。")
    add_step(doc, 6, "登记收付款", "财务在应收管理登记回款，在上游应付登记付款；错误记录通过作废操作纠正并保留日志。")
    add_step(doc, 7, "锁账", "主管在参数规则页查看月度闭环提醒并锁账。未完成项目只作提醒，仍可锁账；锁账后历史月份禁止覆盖导入。")
    add_callout(doc, "月度数据保存方式", "不同月份的数据分别保存。切换月份只改变查询范围，不会删除其他月份；同月份重新导入时，以最新有效批次作为统计来源，旧批次保留审计状态。", fill=LIGHT_GREEN, accent=TEAL)

    add_heading(doc, "4. Excel 导入操作", 1)
    add_heading(doc, "4.1 导入前准备", 2)
    add_list(doc, [
        "使用经营总览中的“上传 Excel 导入”，不要上传截图、PDF 或改名后的临时文件。",
        "工作表必须包含运单号、收付类型、费用类型等核心字段；系统会自动寻找有效明细工作表。",
        "保留原始正负号、币种/汇率标注、供应商、销售代表、客服代表、用户和下单时间。",
        "单个上传文件默认不超过 25MB，仅支持 .xlsx 与 .xls。",
    ])
    add_heading(doc, "4.2 两步导入", 2)
    add_step(doc, 1, "选择文件并等待预检", "预检阶段不写数据库。系统解析表头、月份、明细行、订单数、应收、应付、毛利、物流/服务拆分和质量问题。")
    add_step(doc, 2, "核对字段映射", "确认系统字段与 Excel 原始表头一一对应。缺失必填字段属于阻断问题；额外表头或非关键缺失通常作为提示。")
    add_step(doc, 3, "处理阻断项", "按提示中的 Excel 行号和字段修改原始表格后重新上传。存在阻断项时按钮不可写入数据库。")
    add_step(doc, 4, "确认写入数据库", "确认月份、票数及金额无误后点击“确认写入数据库”。成功后系统自动切换到导入月份并刷新全部页面。")
    add_callout(doc, "严禁直接修正汇总数字", "若系统金额与 Excel 透视表不同，应从订单费用明细和原始 Excel 行核对收付方向、原始金额、汇率、费用类型及重复行。不能通过手工修改总应收或总应付来消除差异。", fill=LIGHT_ORANGE, accent=ORANGE)

    add_heading(doc, "4.3 固定表头规范", 2)
    headers = [
        "运单号", "客户订单号", "用户", "服务", "收费重(KG)", "供应商收费重(KG)", "供应商", "供应商服务",
        "收付类型", "费用类型", "金额", "单价", "本币费用", "销售代表", "备注", "备注", "折合人民币",
        "客服代表", "下单时间", "内部备注", "实重", "件数", "主品名",
    ]
    rows = []
    for index in range(0, len(headers), 4):
        part = headers[index:index + 4]
        rows.append([str(index + offset + 1) if offset < len(part) else "" for offset in range(4)] + part + [""] * (4 - len(part)))
    # Rebuild into four index/header pairs for compact reference.
    matrix = []
    for start in range(0, len(headers), 4):
        chunk = headers[start:start + 4]
        row = []
        for offset in range(4):
            if offset < len(chunk):
                row.extend([str(start + offset + 1), chunk[offset]])
            else:
                row.extend(["", ""])
        matrix.append(row)
    add_table(doc, ["序", "表头", "序", "表头", "序", "表头", "序", "表头"], matrix, [600, 1740, 600, 1740, 600, 1740, 600, 1740])
    doc.add_page_break()

    add_heading(doc, "5. 经营与利润页面", 1)
    add_heading(doc, "5.1 经营总览", 2)
    add_list(doc, [
        "顶部指标：总应收、总应付、调整后毛利、毛利率、总票数、物流提成、高风险票数。",
        "月度趋势：鼠标悬停到月份数据点，查看该月应收、应付、毛利和毛利率。",
        "业务类型利润：比较各业务类型本月应收、毛利与毛利率。",
        "业务员毛利、客户利润、上游供应商应付集中度：用于快速定位重点人员、客户和供应商。",
        "接口加载失败时页面显示错误与重试，不应把失败数据误显示为零。",
    ])
    add_heading(doc, "5.2 业务利润", 2)
    add_list(doc, [
        "总业务卡片 = 物流业务 + 注册/服务类业务。",
        "物流与注册/服务分类分别显示票数、应收、应付、毛利和毛利率。",
        "业务类型汇总必须与订单明细和费用行合计一致；服务类不进入物流提成口径。",
    ])
    add_heading(doc, "5.3 客户利润分析", 2)
    add_list(doc, [
        "应收占比用于识别最大流水客户，毛利占比用于识别利润贡献客户。",
        "客户汇总包含应收、应付、毛利和毛利率；可按应收或毛利排序比较。",
        "客户数据仅统计物流类应收应付时，页面会明确提示口径。",
    ])

    add_heading(doc, "6. 物流提成与注册提成", 1)
    add_heading(doc, "6.1 物流提成", 2)
    add_step(doc, 1, "查看销售代表汇总", "核对票数、物流调整后毛利、汇总毛利率、提成比例和应记提成。")
    add_step(doc, 2, "下钻订单", "点击票数查看系统单号、原始订单号、对应用户、毛利、提成比例、提成金额和状态。")
    add_step(doc, 3, "调整比例", "有权限的主管可修改单票提成比例。保存后系统自动重算提成金额；已确认历史单据不能直接覆盖。")
    add_step(doc, 4, "确认并生成单据", "完成核对后确认销售代表提成，并生成个人确认单。")
    add_heading(doc, "6.2 注册提成", 2)
    add_list(doc, [
        "数据来自 Excel 中被识别为公司注册、证书、商标、店铺租赁等服务类订单。",
        "核对对应用户、销售代表、客服代表、成交单价、成交利润、限定条件和建议提成。",
        "主管填写确认提成并保存；确认金额归属对应销售代表，进入销售薪资确认单。",
        "注册/服务提成不进入物流提成汇总，避免重复计算。",
    ])
    doc.add_page_break()

    add_heading(doc, "7. 操作员绩效", 1)
    add_list(doc, [
        "人员归属严格读取 Excel 的客服代表字段，页面称为“操作员”。",
        "实际票数按导入 Excel 订单统计，不人工增删；规则基础票数按业务类型自动匹配。",
        "空运白关固定按 50 元/票计算；其他业务按页面规则卡和当前参数执行。",
        "允许修改的绩效规则值、分类金额和发放说明保存到数据库，系统自动汇总最终绩效金额。",
        "保存调整不修改 Excel 原始票数和原始台账。",
    ])
    add_callout(doc, "核对重点", "每名操作员的最终绩效金额应等于各绩效板块小计之和。个人确认单会分别展示业务类型、Excel 票数、规则基础票数、计发票数、绩效规则和分类金额。", fill=LIGHT_GREEN, accent=TEAL)

    add_heading(doc, "8. 电子签名确认", 1)
    add_heading(doc, "8.1 生成确认单", 2)
    add_step(doc, 1, "批量生成", "在电子签名确认页生成销售薪资确认单和操作员薪资确认单。销售单汇总物流提成及已确认注册/服务提成；操作员单汇总绩效。")
    add_step(doc, 2, "核对快照", "点击“查看确认单”，核对月份、人员、订单/绩效明细、金额构成、最终金额和发放说明。")
    add_step(doc, 3, "下载文件", "Excel、PDF、PNG 均从同一份确认单快照生成，版本、人员和金额应一一对应。")
    add_heading(doc, "8.2 外发和签名", 2)
    add_step(doc, 1, "生成外发链接", "点击生成外发链接。配置钉钉且已维护员工钉钉用户 ID 时，可直接发送；否则复制完整链接手工发送。")
    add_step(doc, 2, "员工查看并签名", "员工打开公开签名页，无需后台登录即可查看自己的确认单明细，确认无误后提交电子签名。")
    add_step(doc, 3, "主管确认", "员工签名后，主管回到后台核对并执行主管确认。签名时间、IP、浏览器信息和主管确认均写入证据链。")
    add_step(doc, 4, "作废重签", "已主管确认的单据不能覆盖。发现错误时填写作废原因，再生成新版本并重新发送。")
    add_callout(doc, "一次性链接", "员工成功签名后 token 立即失效。链接过期、已签名或单据作废后必须重新生成，不能重复提交。", fill=LIGHT_ORANGE, accent=ORANGE)

    add_heading(doc, "9. 风险复查", 1)
    add_step(doc, 1, "选择风险范围", "可筛选全部风险、高风险或异常高利润订单。")
    add_step(doc, 2, "查看原始数据", "打开原始数据明细，按 Excel 行号核对系统单号、原始订单号、对应用户、收付方向、费用类型、原始金额、本币金额、汇率、供应商和人员归属。")
    add_step(doc, 3, "填写复核说明", "记录核对结论和处理意见后保存。不要只点击完成而不写说明。")
    add_step(doc, 4, "保留审计记录", "复核人、时间、状态和说明会保存到数据库并写入操作日志。")
    doc.add_page_break()

    add_heading(doc, "10. 应收管理与上游应付", 1)
    add_heading(doc, "10.1 应收管理", 2)
    add_list(doc, [
        "查看总应收、已回款、未回款、逾期未回款和 0-30/31-60/61-90/90+ 天账龄。",
        "按客户或订单搜索，核对订单编号、原始订单号、应收、已收、未收及状态。",
        "登记回款时填写金额、日期、方式、参考号和备注；保存后订单状态、汇总和账龄同步刷新。",
        "错误回款使用“作废”纠正，不能删除数据库记录；作废原因写入日志。",
    ])
    add_heading(doc, "10.2 上游应付", 2)
    add_list(doc, [
        "页面只统计物流类订单；注册、证书、店铺租赁等服务类应付单独管理。",
        "供应商汇总按 Excel 原始应付费用行归集；同一订单多个供应商会分别计入对应供应商。",
        "查看总应付、已付款、未付款、逾期未付款和账龄，并下钻订单费用明细。",
        "登记付款或作废错误付款后，供应商汇总、订单状态、未付款和账龄同步更新。",
    ])
    add_table(
        doc,
        ["核对关系", "正确结果"],
        [
            ["订单应收合计", "等于应收管理总应收及经营总览总应收"],
            ["物流订单应付合计", "等于上游应付总应付"],
            ["应收 - 应付", "等于对应范围的调整后毛利"],
            ["已收/已付 + 未收/未付", "等于订单应收/应付"],
        ],
        [3000, 6360],
    )

    add_heading(doc, "11. 参数规则与月结锁账", 1)
    add_heading(doc, "11.1 月度闭环", 2)
    add_list(doc, [
        "流程步骤包含 Excel 导入、导入审计、风险复核、注册确认、提成签名、操作员签名、应收应付、CFO 就绪和锁账。",
        "未完成事项显示为提醒，不阻止主管锁账；锁账前应根据实际业务判断是否允许带提醒关闭。",
        "锁账后禁止覆盖导入、回滚批次、重新生成确认单和修改该月历史数据。需要调整时由主管解锁并填写原因。",
        "锁账和解锁均写入操作日志。",
    ])
    add_heading(doc, "11.2 参数与模板", 2)
    add_list(doc, [
        "表头模板只保存固定表头规范，不保存业务数据。后续 Excel 导入按模板进行字段映射。",
        "参数规则使用 JSON 保存。修改前先备份，保存后影响后续导入和计算，不应追溯重算已确认历史快照。",
        "只有具有 rules:write 权限的管理员可以保存参数规则。",
    ])
    add_heading(doc, "11.3 账号与操作日志", 2)
    add_list(doc, [
        "管理员可创建账号、停用账号、重置密码和维护员工钉钉用户 ID。",
        "不要多人共用管理员账号；每名员工使用独立账号，便于追溯操作人。",
        "导入、回滚、风险复核、收付款、锁账、规则修改、签名和主管确认均可在操作日志中查询。",
    ])
    doc.add_page_break()

    add_heading(doc, "12. 系统运行状态与故障处理", 1)
    add_heading(doc, "12.1 运维状态面板", 2)
    add_text(doc, "参数规则页的“系统运行与就绪状态”用于判断问题发生在前端、后端、数据库还是外部集成。")
    add_table(
        doc,
        ["指标", "正常状态", "异常处理"],
        [
            ["数据库", "通过，延迟为正常毫秒值", "检查 Render PostgreSQL 或本地 54329"],
            ["表头模板/参数规则", "通过", "刷新模板或检查数据库初始化"],
            ["请求失败/慢请求", "失败为 0，P95 合理", "按请求 ID 查询后端结构化日志"],
            ["进程内存", "稳定，无持续增长", "检查大文件导入、循环请求和 Render 实例"],
            ["钉钉/ERPNext", "已配置或明确显示未配置", "检查 Render 环境变量和第三方权限"],
        ],
        [2200, 2600, 4560],
    )
    add_heading(doc, "12.2 常见问题", 2)
    add_table(
        doc,
        ["问题", "判断方法", "处理方式"],
        [
            ["网页打不开", "区分本地地址还是线上地址", "本地启动前后端和数据库；线上检查 Pages/Render"],
            ["登录失败", "查看是否账号停用、密码已改或 token 失效", "联系管理员重置；重新登录，不反复尝试默认密码"],
            ["导入失败", "查看预检阻断项、文件类型、大小及月份锁账状态", "修复指定 Excel 行/字段后重新预检"],
            ["金额与 Excel 不一致", "下钻费用行，核对方向、金额、汇率、重复行", "修正源 Excel 后重新导入；不要改汇总"],
            ["下载 401/403", "确认登录状态和角色权限", "重新登录后使用页面下载按钮，不直接打开 API URL"],
            ["签名链接打不开", "确认使用线上链接、token 未过期且单据未作废", "重新生成并发送外部签名链接"],
            ["钉钉发送失败", "查看钉钉配置、应用发布和员工 ID 映射", "修复配置后重发，或复制链接手工发送"],
            ["Render 首次访问慢", "健康接口等待时间明显增加", "免费实例可能冷启动；升级实例可减少休眠"],
        ],
        [1900, 3140, 4320],
    )
    add_callout(doc, "请求 ID", "后端为每次请求返回 x-request-id。出现接口错误时记录页面时间、操作、月份和请求 ID，便于在 Render 日志中定位同一条请求。", fill=LIGHTER_BLUE, accent=BLUE)

    add_heading(doc, "13. 钉钉与 ERPNext", 1)
    add_heading(doc, "13.1 钉钉签名通知", 2)
    add_list(doc, [
        "企业应用单聊需要配置 AppKey、AppSecret、RobotCode，授予发送消息权限并发布应用。",
        "参数规则页为每名员工维护正确的钉钉用户 ID。",
        "先用测试员工生成确认单并发送；成功后再批量使用。",
        "群机器人 Webhook 可作为群通知方式，但个人签名仍应使用每人独立链接。",
    ])
    add_heading(doc, "13.2 ERPNext", 2)
    add_list(doc, [
        "当前集成为只读模式，通过后端安全读取 ERPNext 客户、供应商和销售/采购发票概览。",
        "ERPNext 原始金额和币种保持不变，不自动写入 XJD 财务台账。",
        "先在页面点击“测试连接”，成功后再查看概览或打开 ERPNext。",
    ])

    add_heading(doc, "14. 报表、下载与备份", 1)
    add_list(doc, [
        "经营总览“导出月报”输出当前月份财务分析 Excel。",
        "确认单 Excel、PDF、PNG 使用同一份 payload 快照，历史版本不重新计算金额。",
        "参数规则页可导出当前月份或全量系统备份，包含模板、规则、批次、锁账、确认单和操作日志。",
        "所有后台下载均通过当前登录 token 获取；不要直接转发受保护的下载 API。",
    ])
    doc.add_page_break()

    add_heading(doc, "15. 月度操作检查清单", 1)
    add_heading(doc, "15.1 导入前", 2)
    add_list(doc, [
        "确认目标月份及该月是否锁账。",
        "确认 Excel 使用固定表头，明细工作表完整。",
        "确认销售代表、客服代表、用户、供应商和下单时间已填写。",
        "确认金额正负号、币种/汇率标注和赔付费用未被改写。",
    ])
    add_heading(doc, "15.2 导入后", 2)
    add_list(doc, [
        "核对应收、应付、毛利、票数和物流/服务拆分。",
        "核对费用明细合计与订单及汇总差异为 0。",
        "完成风险复查、注册提成、物流提成和操作员绩效确认。",
        "生成并发送确认单，跟进员工签名和主管确认。",
        "登记本月回款和付款，处理错误结算记录。",
    ])
    add_heading(doc, "15.3 锁账前", 2)
    add_list(doc, [
        "查看月度闭环提醒并确认未完成事项是否允许带入后续处理。",
        "导出月报和系统备份。",
        "填写锁账原因并执行锁账；确认该月份后续不能直接覆盖导入。",
        "抽查操作日志、确认单版本和签名证据。",
    ])
    add_callout(doc, "最终责任", "系统提供计算、追溯和流程留痕，但不替代正式会计凭证、银行流水、合同、发票和主管审批。财务结论应由具备权限的财务或主管人员最终确认。", fill=LIGHT_ORANGE, accent=ORANGE)

    add_heading(doc, "附录：支持与版本", 1)
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["系统名称", "XJD Finance 跨境物流财务管理"],
            ["前端部署", "GitHub Pages"],
            ["后端与数据库", "Render Web Service + PostgreSQL"],
            ["本地端口", "前端 5173；后端 4000；PostgreSQL 54329"],
            ["健康检查", "/api/health；数据库就绪 /api/health/ready"],
            ["手册版本", f"1.0（{date.today().isoformat()}）"],
        ],
        [2500, 6860],
    )

    doc.core_properties.title = "XJD Finance 系统使用操作说明书"
    doc.core_properties.subject = "跨境物流财务管理系统操作指南"
    doc.core_properties.author = "XJD Finance"
    doc.core_properties.keywords = "XJD Finance, 财务管理, Excel导入, 提成, 电子签名, 应收应付"
    doc.core_properties.comments = "根据当前系统功能与权限生成。"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_manual()
